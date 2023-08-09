# TODO: trial inputting charts into table
# also remove cfd files from project -> move
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches

graph_path = 'data\escape_prob_chart.png'
document_name = 'Event Tree.docx'
document_path = Path(__file__).parent /"Report Template"/"Table Images.docx"
# doc = DocxTemplate(document_path)
# output_path = Path(__file__).parent /"Report Template"/"Table Trial.docx"
document = Document(document_path)
# not using template but access table cells
pass
table = [t for t in document.tables][-1]
document.add_picture('logo.png', width=Inches(2), height=Inches(2))

def create_inline_image(image_file, template=document):
    return InlineImage(template, image_descriptor=image_file, width=Inches(6), height=Inches(4))
